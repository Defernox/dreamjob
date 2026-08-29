"""Manipulation d'un .docx **sans jamais casser sa mise en page**.

Principe unique : on ne crée jamais un paragraphe de zéro. On part toujours d'un
paragraphe existant du modèle — dont on connaît le style, la police, les marges
et les puces — et on le duplique ou on remplace son texte. Le modèle reste donc
la seule source de vérité de l'apparence.
"""

from __future__ import annotations

import copy

from docx.text.paragraph import Paragraph


def definir_texte(paragraphe: Paragraph, texte: str) -> None:
    """Remplace le contenu en gardant la mise en forme du premier run."""
    if not paragraphe.runs:
        paragraphe.add_run(texte)
        return
    paragraphe.runs[0].text = texte
    for run in paragraphe.runs[1:]:
        run._element.getparent().remove(run._element)


def definir_morceaux(paragraphe: Paragraph, morceaux: list[str]) -> None:
    """Remplit un paragraphe à plusieurs runs (ex. « Catégorie : » en gras puis
    le contenu en normal), en conservant la mise en forme de chaque run."""
    if not paragraphe.runs:
        paragraphe.add_run("".join(morceaux))
        return

    for i, run in enumerate(list(paragraphe.runs)):
        if i < len(morceaux):
            run.text = morceaux[i]
        else:
            run._element.getparent().remove(run._element)

    # Plus de morceaux que de runs : le surplus rejoint le dernier run.
    if len(morceaux) > len(paragraphe.runs):
        paragraphe.runs[-1].text += "".join(morceaux[len(paragraphe.runs):])


def cloner_apres(modele: Paragraph, reference: Paragraph) -> Paragraph:
    """Copie `modele` (mise en forme comprise) et l'insère après `reference`."""
    nouveau = copy.deepcopy(modele._p)
    reference._p.addnext(nouveau)
    return Paragraph(nouveau, reference._parent)


def copier_xml(paragraphes: list[Paragraph]) -> list:
    """Photographie d'un bloc AVANT qu'on le remplisse.

    Sans cela, dupliquer un bloc déjà rempli recopierait le texte du premier
    élément au lieu de la mise en forme vierge du modèle.
    """
    return [copy.deepcopy(p._p) for p in paragraphes]


def cloner_xml_apres(element, reference: Paragraph) -> Paragraph:
    nouveau = copy.deepcopy(element)
    reference._p.addnext(nouveau)
    return Paragraph(nouveau, reference._parent)


def supprimer(paragraphe: Paragraph) -> None:
    element = paragraphe._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def est_titre(paragraphe: Paragraph) -> bool:
    style = paragraphe.style
    return style is not None and style.name.startswith("Heading")


def est_puce(paragraphe: Paragraph) -> bool:
    style = paragraphe.style
    return style is not None and style.name == "List Paragraph"


def est_gras(paragraphe: Paragraph) -> bool:
    return any(run.bold for run in paragraphe.runs if run.bold)


def decouper_en_sections(document) -> tuple[list[Paragraph], dict[str, list[Paragraph]]]:
    """Renvoie (paragraphes d'en-tête, {titre de section: paragraphes}).

    Les titres servent de bornes : le modèle peut être réorganisé ou traduit
    sans toucher au code, tant que ses intitulés restent reconnaissables.
    """
    entete: list[Paragraph] = []
    sections: dict[str, list[Paragraph]] = {}
    courante: str | None = None

    for paragraphe in document.paragraphs:
        if est_titre(paragraphe):
            courante = paragraphe.text.strip()
            sections[courante] = []
        elif courante is None:
            entete.append(paragraphe)
        else:
            sections[courante].append(paragraphe)
    return entete, sections


def titre_de_section(document, nom: str) -> Paragraph | None:
    for paragraphe in document.paragraphs:
        if est_titre(paragraphe) and paragraphe.text.strip().lower() == nom.lower():
            return paragraphe
    return None


def supprimer_section(document, nom: str) -> None:
    """Retire un titre de section et tout son contenu — pour ne pas laisser de
    rubrique vide (« Certifications » sans certification) dans le CV rendu."""
    titre = titre_de_section(document, nom)
    if titre is None:
        return
    _, sections = decouper_en_sections(document)
    for paragraphe in sections.get(titre.text.strip(), []):
        supprimer(paragraphe)
    supprimer(titre)
