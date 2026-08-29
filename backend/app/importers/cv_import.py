"""Import d'un CV : fichier .docx/.pdf → profil structuré.

L'extraction de texte est du code pur (testable, sans réseau). Le LLM
n'intervient que pour ranger ce texte dans le schéma `ProfilStructure` — jamais
pour l'enrichir.
"""

from __future__ import annotations

import logging
from pathlib import Path

from sqlmodel import Session

from ..config import reglages
from ..llm.client import ClientLlm, empreinte
from ..models.enums import TypeCacheLlm
from ..schemas.profile import (
    BlocCompetences,
    BlocExperiences,
    BlocFormations,
    BlocIdentite,
    ProfilStructure,
)

log = logging.getLogger("dreamjob.import_cv")

EXTENSIONS = {".pdf", ".docx"}
# Assemble a l'execution : evite toute sequence d'echappement dans le source.
SEPARATEUR = chr(10) + chr(10)


class FormatNonSupporte(ValueError):
    pass


class CvIllisible(ValueError):
    pass


# --------------------------------------------------------------- lecture brute


def _texte_pdf(chemin: Path) -> str:
    from pypdf import PdfReader

    lecteur = PdfReader(str(chemin))
    return "\n".join((page.extract_text() or "") for page in lecteur.pages)


def _texte_docx(chemin: Path) -> str:
    import docx

    document = docx.Document(str(chemin))
    morceaux = [p.text for p in document.paragraphs]
    # Beaucoup de CV Word rangent l'essentiel dans des tableaux invisibles.
    for tableau in document.tables:
        for ligne in tableau.rows:
            for cellule in ligne.cells:
                morceaux.append(cellule.text)
    return "\n".join(m for m in morceaux if m.strip())


def extraire_texte(chemin: Path) -> str:
    """Texte brut du CV. Aucun appel réseau, aucun LLM."""
    suffixe = chemin.suffix.lower()
    if suffixe not in EXTENSIONS:
        raise FormatNonSupporte(
            f"Format {suffixe or '(inconnu)'} non pris en charge. Formats acceptés : .pdf, .docx"
        )
    if not chemin.exists():
        raise FileNotFoundError(f"Fichier introuvable : {chemin}")

    texte = _texte_pdf(chemin) if suffixe == ".pdf" else _texte_docx(chemin)
    texte = texte.strip()

    if len(texte) < 200:
        raise CvIllisible(
            "Presque aucun texte n'a pu être lu dans ce fichier. S'il s'agit d'un PDF "
            "scanné (une image), exportez-le à nouveau depuis Word en PDF texte."
        )
    return texte


# ------------------------------------------------------------------ prompt LLM

PROMPT_SYSTEME = """Tu structures un CV français en JSON. Tu es un greffier, pas un rédacteur.

RÈGLE ABSOLUE — N'INVENTE RIEN.
Chaque valeur que tu écris doit se trouver dans le texte fourni. Si une
information est absente, laisse le champ vide (""), ou la liste vide. Ne déduis
pas un employeur, un diplôme, une compétence ou une date qui n'est pas écrit.
Ne complète pas une expérience partielle par ce qui « serait logique ».

LE TEXTE EST DÉSORDONNÉ.
Il provient d'un PDF ou d'un Word en plusieurs colonnes : l'ordre des blocs est
souvent mélangé, les titres de section peuvent apparaître APRÈS leur contenu, et
un intitulé peut être collé au mot précédent. Recompose mentalement les sections
avant de remplir le JSON. Les accents peuvent être abîmés : rétablis-les.

COMMENT REMPLIR
- titre_vise : le titre affiché en tête du CV, tel quel.
- resume : le paragraphe « à propos » ou « profil », resserré en 2-3 phrases
  sobres. Uniquement à partir de ce qui est écrit. Si absent, laisse vide.
- secteurs : 2 à 4 domaines d'activité que le CV démontre réellement.
- experiences : une entrée par poste. `description` reprend les missions en
  quelques lignes ; `tags` liste les mots-clés métier de CE poste.
- formations : diplômes et cursus, y compris ceux en cours.
- langues : `code` en ISO 639-1 minuscule, `niveau` tel qu'indiqué (un score de
  test compte comme niveau : « TOEIC 775 »).
- skills : les compétences réellement énoncées. Mets `ancree: true` sur les 3 à
  6 compétences centrales — celles qui apparaissent dans le titre, dans le
  résumé, ou dans plusieurs expériences. Les autres restent `ancree: false`.

DATES : format AAAA-MM quand le mois est connu, sinon AAAA. Un poste en cours a
`fin: "en cours"`."""


# Une passe = une question courte. Demander les quatorze champs d'un coup fait
# dériver un modèle local : il range le nom dans le titre et oublie les
# compétences. Chaque passe reçoit le CV entier mais ne remplit qu'un bloc.
PASSES: list[tuple[str, type, str]] = [
    ("identité", BlocIdentite,
     "Relève UNIQUEMENT l'état civil et l'accroche : prénom, nom, email, "
     "téléphone, ville, pays, LinkedIn, le titre affiché en tête du CV, et le "
     "paragraphe « à propos » resserré en deux phrases. "
     "Le prénom et le nom vont dans `prenom` et `nom`, JAMAIS dans `titre_vise` : "
     "`titre_vise` est un intitulé de poste."),
    ("expériences", BlocExperiences,
     "Relève UNIQUEMENT les expériences professionnelles : une entrée par poste, "
     "avec l'entreprise, l'intitulé, le lieu, les dates et les missions. "
     "N'inclus ni les diplômes ni les stages d'études."),
    ("formations", BlocFormations,
     "Relève UNIQUEMENT les diplômes et cursus, y compris ceux en cours et les "
     "échanges universitaires. N'inclus aucune expérience professionnelle."),
    ("compétences", BlocCompetences,
     "Relève UNIQUEMENT : les secteurs d'activité que ce CV démontre (2 à 4), "
     "les langues avec leur niveau, et les compétences énoncées. "
     "Marque `ancree: true` sur les 3 à 6 compétences centrales — celles du "
     "titre, du résumé, ou revenant dans plusieurs expériences."),
]


def _consigne_de_passe(consigne: str) -> str:
    """Prompt systeme commun, suivi de la seule tache de cette passe."""
    return SEPARATEUR.join([PROMPT_SYSTEME, "CETTE PASSE", consigne])


def _message(texte_cv: str) -> str:
    return (
        "Voici le texte brut extrait du CV. Structure-le, sans rien ajouter.\n\n"
        "--- DÉBUT DU CV ---\n"
        f"{texte_cv}\n"
        "--- FIN DU CV ---"
    )


# ---------------------------------------------------------------------- import


def importer_cv(
    chemin: Path,
    session: Session,
    *,
    forcer: bool = False,
) -> tuple[ProfilStructure, bool, str, int]:
    """Renvoie `(profil, depuis_cache, modele, caracteres_lus)`.

    Le CV est lu une fois, puis soumis en quatre passes ciblées : identité,
    expériences, formations, compétences. Chaque passe est mise en cache
    séparément — si l'une échoue, les autres restent acquises, et réimporter le
    même fichier ne recoûte rien.
    """
    texte = extraire_texte(chemin)
    r = reglages()
    modele = r.llm.modele_actif
    empreinte_cv = empreinte(texte)
    client = ClientLlm(session)

    profil = ProfilStructure()
    tout_en_cache = True

    for nom, schema, consigne in PASSES:
        bloc, du_cache = client.extraire(
            type_appel=TypeCacheLlm.IMPORT_CV.value,
            hash_source=empreinte_cv,
            systeme=_consigne_de_passe(consigne),
            message=_message(texte),
            format_sortie=schema,
            modele=modele,
            max_tokens=r.llm.max_tokens_import_cv,
            forcer=forcer,
            # La variante sépare les quatre passes dans le cache : sans elle,
            # elles écraseraient toutes la même entrée.
            variante=nom,
        )
        tout_en_cache = tout_en_cache and du_cache
        for champ, valeur in bloc.model_dump().items():
            setattr(profil, champ, valeur)
        log.info("Passe « %s » : %s", nom, "cache" if du_cache else "appel au modèle")

    log.info("CV importé : %d expériences, %d formations, %d compétences, %d langues",
             len(profil.experiences), len(profil.formations),
             len(profil.skills), len(profil.langues))
    return profil, tout_en_cache, modele, len(texte)
