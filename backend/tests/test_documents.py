"""Rendu du CV et fabrication du dossier de candidature.

Exigence du projet : « La mise en page du modèle ne doit jamais être cassée. »
On vérifie donc que les styles du modèle survivent au rendu.
"""

import json
from pathlib import Path

import docx
import pytest

from app.documents.cv_render import ModeleIntrouvable, rendre
from app.documents.dossier import ecrire_lettre, generer, nom_dossier, slug
from app.models import Offer, Profile

MODELE = Path(__file__).resolve().parents[2] / "templates" / "cv_modele.docx"

pytestmark = pytest.mark.skipif(
    not MODELE.exists(), reason="templates/cv_modele.docx absent"
)


@pytest.fixture
def profil():
    return Profile(
        prenom="Maxime", nom="Nicolas", ville="Paris", pays="France",
        email="test@example.com", telephone="06 00 00 00 00",
        titre_vise="Analyste financier", resume="Parcours en finance de marché.",
        secteurs=["banque"], pays_acceptes=["France"], contrats_acceptes=["CDI"],
        langues=[{"code": "fr", "libelle": "Français", "niveau": "natif"}],
        skills=[{"nom": "Analyse financière", "ancree": True},
                {"nom": "Recouvrement de créances", "ancree": False},
                {"nom": "Excel", "ancree": False}],
        experiences=[
            {"entreprise": "Association Sport", "poste": "Trésorier", "debut": "2021-09",
             "fin": "2022-09", "description": "Tenue de la trésorerie. Budget annuel."},
            {"entreprise": "Crédit Mutuel", "poste": "Gestionnaire export",
             "debut": "2023-09", "fin": "2025-09",
             "description": "Recouvrement de créances export. Analyse financière des dossiers."},
        ],
        formations=[{"etablissement": "EM Normandie", "diplome": "Master 2 Finance",
                     "annee": "2020-2025", "details": "Marchés financiers."},
                    {"etablissement": "ESLSCA", "diplome": "MBA Trading", "annee": "2026"}],
    )


@pytest.fixture
def offre():
    return Offer(source="test", source_id="1", titre="Chargé de recouvrement de créances",
                 entreprise="Assureur Crédit", lieu="92 - Nanterre", pays="France",
                 type_contrat="CDI",
                 description_brute="Recouvrement de créances export et analyse financière.")


def _textes(chemin: Path) -> list[str]:
    return [p.text.strip() for p in docx.Document(str(chemin)).paragraphs if p.text.strip()]


# --- Rendu du CV ------------------------------------------------------------


def test_le_cv_contient_tout_le_profil(profil, offre, tmp_path):
    joint = " | ".join(_textes(rendre(profil, offre, MODELE, tmp_path / "CV.docx")))
    assert "MAXIME NICOLAS" in joint
    for experience in profil.experiences:
        assert experience["entreprise"] in joint
    for formation in profil.formations:
        assert formation["diplome"] in joint


def test_aucun_crochet_du_modele_ne_subsiste(profil, offre, tmp_path):
    """Les [placeholders] du modèle doivent tous avoir été remplacés."""
    for texte in _textes(rendre(profil, offre, MODELE, tmp_path / "CV.docx")):
        assert "[" not in texte, f"placeholder non remplacé : {texte[:60]}"


def test_la_mise_en_page_du_modele_survit(profil, offre, tmp_path):
    """Les styles du modèle SONT la mise en page : ils doivent rester."""
    document = docx.Document(str(rendre(profil, offre, MODELE, tmp_path / "CV.docx")))
    styles = {p.style.name for p in document.paragraphs if p.style is not None}
    assert "Heading 2" in styles, "les titres de section ont perdu leur style"
    assert "List Paragraph" in styles, "les puces ont perdu leur style"


def test_les_experiences_sont_reordonnees_selon_l_offre(profil, offre, tmp_path):
    """L'expérience la plus proche de l'annonce passe en tête."""
    textes = _textes(rendre(profil, offre, MODELE, tmp_path / "CV.docx", reordonner=True))
    position = {t: i for i, t in enumerate(textes)}
    assert position["Gestionnaire export"] < position["Trésorier"]


def test_l_ordre_du_profil_est_respecte_si_on_desactive(profil, offre, tmp_path):
    textes = _textes(rendre(profil, offre, MODELE, tmp_path / "CV.docx", reordonner=False))
    position = {t: i for i, t in enumerate(textes)}
    assert position["Trésorier"] < position["Gestionnaire export"]


def test_les_puces_restent_dans_leur_experience(profil, offre, tmp_path):
    """Bug corrigé : les puces d'un bloc migraient vers le bloc suivant."""
    document = docx.Document(str(rendre(profil, offre, MODELE, tmp_path / "CV.docx")))
    bloc, puces = None, {}
    for paragraphe in document.paragraphs:
        texte = paragraphe.text.strip()
        if not texte:
            continue
        if paragraphe.style is not None and paragraphe.style.name == "List Paragraph":
            puces.setdefault(bloc, []).append(texte)
        elif any(r.bold for r in paragraphe.runs if r.bold):
            bloc = texte

    assert any("créances export" in p for p in puces.get("Gestionnaire export", []))
    assert any("trésorerie" in p.lower() for p in puces.get("Trésorier", []))


def test_les_rubriques_sans_contenu_disparaissent(profil, offre, tmp_path):
    """Un CV ne doit pas afficher « Certifications » quand il n'y en a aucune."""
    textes = _textes(rendre(profil, offre, MODELE, tmp_path / "CV.docx"))
    assert "Certifications" not in textes
    assert "Projets" not in textes


def test_le_titre_du_cv_reprend_l_intitule_de_l_offre(profil, offre, tmp_path):
    """C'est ce que lisent les filtres automatiques des recruteurs."""
    assert offre.titre in _textes(rendre(profil, offre, MODELE, tmp_path / "CV.docx"))


def test_modele_absent_donne_un_message_utile(profil, offre, tmp_path):
    with pytest.raises(ModeleIntrouvable, match="Déposez votre CV"):
        rendre(profil, offre, tmp_path / "fantome.docx", tmp_path / "CV.docx")


# --- Dossier de candidature -------------------------------------------------


def _redacteur_honnete(systeme, message):
    return "Mon parcours au Crédit Mutuel me prépare à ce poste. " * 25


def test_nom_de_dossier_lisible_et_sans_accent(offre):
    nom = nom_dossier(offre)
    assert nom.startswith("20")
    assert "assureur-credit" in nom
    # Utilisable tel quel sur le disque : ni accent, ni espace, ni ponctuation.
    assert slug(nom, len(nom)) == nom


def test_un_intitule_a_rallonge_ne_produit_pas_un_chemin_trop_long(profil, tmp_path):
    """Windows refuse au-delà de 260 caractères de chemin complet."""
    offre = Offer(source="t", source_id="1",
                  entreprise="Groupe International de Services Financiers et Assurantiels",
                  titre="Chargé de recouvrement de créances export à l'international "
                        "pour la zone Europe Moyen-Orient Afrique (H/F)")
    nom = nom_dossier(offre)
    assert len(nom) <= 80
    assert not nom.endswith("-")


def test_le_dossier_complet_est_produit(profil, offre, tmp_path):
    resultat = generer(profil, offre, tmp_path, MODELE,
                       redacteur=_redacteur_honnete, ouvrir_apres=False)
    noms = {f.name for f in resultat.fichiers}
    assert "CV.docx" in noms
    assert "Lettre_de_motivation.docx" in noms
    assert "offre.json" in noms


def test_l_offre_est_archivee_telle_quelle(profil, offre, tmp_path):
    """L'annonce disparaîtra du site : il faut en garder une copie."""
    offre.score = 82.4
    resultat = generer(profil, offre, tmp_path, MODELE,
                       redacteur=_redacteur_honnete, ouvrir_apres=False)
    archive = json.loads((resultat.dossier / "offre.json").read_text(encoding="utf-8"))
    assert archive["titre"] == offre.titre
    assert archive["description_brute"] == offre.description_brute
    assert archive["score"] == 82.4


def test_une_lettre_refusee_ne_fait_pas_perdre_le_cv(profil, offre, tmp_path):
    """Un CV sans lettre reste utile ; un dossier vide ne l'est pas."""
    def menteur(systeme, message):
        return "J'ai dirigé le fonds Bridgewater pendant dix ans. " * 25

    resultat = generer(profil, offre, tmp_path, MODELE, redacteur=menteur,
                       tentatives_lettre=1, ouvrir_apres=False)
    noms = {f.name for f in resultat.fichiers}
    assert "CV.docx" in noms
    assert "Lettre_de_motivation.docx" not in noms
    assert any("Lettre non générée" in a for a in resultat.avertissements)


def test_la_lettre_est_mise_en_page_a_la_francaise(profil, offre, tmp_path):
    textes = _textes(ecrire_lettre(profil, offre, "Mon parcours.", tmp_path / "L.docx"))
    assert any(t.startswith("Objet : candidature") for t in textes)
    assert "Madame, Monsieur," in textes
    assert any("salutations distinguées" in t for t in textes)
    assert "Assureur Crédit" in textes


# --- Le classement du CV parle la même langue que le score ------------------


def test_le_classement_du_cv_reconnait_les_synonymes():
    """`_pertinence` était une troisième implémentation de l'appariement, sans
    les synonymes : une expérience « risques de crédit » ne rencontrait jamais
    une offre parlant de « credit risk »."""
    from app.documents.cv_render import _pertinence
    from app.scoring.texte import mots

    offre_en = set(mots("credit risk analysis for banking counterparties"))
    assert _pertinence("gestion des risques de crédit", offre_en) > 0.0


def test_le_classement_du_cv_pondere_les_mots_generiques():
    """Une expérience reconnue sur le seul mot « gestion » ne doit pas passer
    devant une expérience réellement pertinente."""
    from app.documents.cv_render import _pertinence
    from app.scoring.texte import mots

    offre = set(mots("gestion des stocks en entrepôt"))
    generique = _pertinence("gestion de trésorerie", offre)
    assert generique < 0.5


def test_le_cv_et_le_score_utilisent_la_meme_mesure():
    """Deux mesures différentes finiraient par se contredire sous les yeux de
    l'utilisateur : un CV qui met en avant ce que le score juge hors sujet."""
    from app.documents.cv_render import _pertinence
    from app.scoring.score import presence
    from app.scoring.texte import mots

    vocabulaire = set(mots("analyse financière et suivi des encours clients"))
    for terme in ("analyse financière", "gestion de trésorerie", "soudure à l'arc"):
        assert _pertinence(terme, vocabulaire) == presence(terme, vocabulaire, flou=False)


# --- Le CV tient sur une page ------------------------------------------------


def test_la_mobilite_ne_liste_pas_tous_les_pays(profil, offre):
    """Dix-sept pays mangeaient trois lignes de l'en-tête — assez pour faire
    déborder le CV sur une seconde page, et personne ne les lit."""
    from app.documents.cv_render import MAX_PAYS_AFFICHES, _mobilite

    profil.pays_acceptes = [f"Pays{i}" for i in range(17)]
    mobilite = _mobilite(profil, offre)
    assert "14 autres pays" in mobilite
    assert mobilite.count(",") < MAX_PAYS_AFFICHES + 1


def test_le_pays_de_l_offre_passe_en_tete(profil, offre):
    """C'est le seul qui intéresse ce recruteur-là."""
    from app.documents.cv_render import _mobilite

    profil.pays_acceptes = ["France", "Belgique", "Luxembourg", "Irlande"]
    offre.pays = "Luxembourg"
    assert _mobilite(profil, offre).startswith("Luxembourg")


def test_peu_de_pays_sont_listes_tels_quels(profil, offre):
    from app.documents.cv_render import _mobilite

    profil.pays_acceptes = ["France", "Belgique"]
    assert _mobilite(profil, offre) == "France, Belgique"


def test_le_plafond_de_puces_est_respecte(profil, offre, tmp_path):
    """`dossier.py` le resserre quand le PDF déborde."""
    from app.documents.cv_render import rendre

    profil.experiences = [{
        "entreprise": "Crédit Mutuel", "poste": "Gestionnaire",
        "debut": "2023", "fin": "2025",
        "description": "\n".join(f"Mission numéro {i} sur le portefeuille." for i in range(8)),
    }]
    cv = rendre(profil, offre, MODELE, tmp_path / "CV.docx", max_puces=2)
    puces = [t for t in _textes(cv) if t.startswith("Mission numéro")]
    assert len(puces) == 2


def test_un_cv_qui_deborde_est_resserre_puis_signale(profil, offre, tmp_path, monkeypatch):
    """Le filet : on mesure le PDF rendu, on ne devine pas la hauteur. Une
    estimation par nombre de caractères se trompait de huit lignes et rabotait
    les expériences pour un débordement imaginaire."""
    from app.documents import dossier as mod

    essais: list[int] = []

    def faux_convertir(source):
        # On feint deux pages tant que les puces ne sont pas descendues à 2.
        pdf = source.with_suffix(".pdf")
        pages = 1 if len(essais) >= 2 else 2
        pdf.write_bytes(b"%PDF-1.4\n" + b"/Type /Page \n" * pages)
        essais.append(pages)
        return pdf

    monkeypatch.setattr(mod.pdf_outil, "convertir", faux_convertir)
    cv, pdf, pages = mod._cv_sur_une_page(profil, offre, MODELE, tmp_path, True)
    assert pages == 1
    assert len(essais) == 3            # il a fallu resserrer deux fois
    assert cv.exists() and pdf is not None


def test_sans_libreoffice_le_cv_part_entier(profil, offre, tmp_path, monkeypatch):
    """Faute de pouvoir mesurer, mieux vaut un CV complet qu'un CV amputé au
    hasard."""
    from app.documents import dossier as mod

    def refuse(source):
        raise RuntimeError("LibreOffice absent")

    monkeypatch.setattr(mod.pdf_outil, "convertir", refuse)
    cv, pdf, pages = mod._cv_sur_une_page(profil, offre, MODELE, tmp_path, True)
    assert cv.exists()
    assert pdf is None and pages == 0
