"""Correctifs issus de la revue des modules documents/ et exports/.

Ces défauts touchaient tous des fichiers que le candidat envoie : une lettre
amputée, un PDF périmé, un CV qui se contredit.
"""

from datetime import date
from pathlib import Path

import docx
import pytest

from app.documents.lettre import nettoyer
from app.exports.excel import _en_date

MODELE = Path(__file__).resolve().parents[2] / "templates" / "cv_modele.docx"
avec_modele = pytest.mark.skipif(not MODELE.exists(), reason="cv_modele.docx absent")


# --- La lettre n'est plus amputée -------------------------------------------


def test_une_formule_en_milieu_de_lettre_ne_coupe_rien():
    """« Dans l'attente… » ouvre couramment un paragraphe de milieu de lettre.
    L'ancienne version supprimait tout à partir de là : 25 mots sur 34."""
    lettre = ("Premier paragraphe qui présente ma candidature au poste visé.\n\n"
              "Dans l'attente de pouvoir en discuter, je précise que mon expérience "
              "au Crédit Mutuel couvre ce périmètre.\n\n"
              "Troisième paragraphe qui détaille mes compétences techniques.")
    assert len(nettoyer(lettre).split()) == len(lettre.split())
    assert "Troisième paragraphe" in nettoyer(lettre)


def test_une_formule_finale_est_bien_retiree():
    lettre = ("Mon parcours au Crédit Mutuel me prépare à ce poste.\n\n"
              "Cordialement,\nMaxime Nicolas")
    resultat = nettoyer(lettre)
    assert "Cordialement" not in resultat
    assert "Crédit Mutuel" in resultat


def test_une_formule_collee_au_dernier_paragraphe_est_retiree():
    lettre = "Mon parcours me prépare à ce poste.\nCordialement, Maxime Nicolas"
    assert "Cordialement" not in nettoyer(lettre)


def test_un_objet_produit_par_le_modele_est_retire():
    """Le document en pose un lui-même : les deux se retrouveraient l'un sous
    l'autre sur un courrier envoyé à un recruteur."""
    lettre = ("Objet : candidature au poste d'analyste\n\nMadame, Monsieur,\n\n"
              "Mon parcours au Crédit Mutuel me prépare à ce poste.")
    resultat = nettoyer(lettre)
    assert "Objet" not in resultat
    assert "Madame" not in resultat
    assert resultat.startswith("Mon parcours")


def test_une_lettre_aux_paragraphes_courts_reste_entiere():
    """Une heuristique « paragraphe court = signature » mangeait les fins de
    lettre laconiques."""
    lettre = "Premier paragraphe.\n\nDeuxième paragraphe.\n\nTroisième paragraphe."
    assert nettoyer(lettre) == lettre


# --- Les dates saisies à la main sont lues ----------------------------------


@pytest.mark.parametrize("saisie", ["15/09/2026", "2026-09-15", "15-09-2026", "15.09.2026"])
def test_une_date_tapee_a_la_main_est_comprise(saisie):
    assert _en_date(saisie) == date(2026, 9, 15)


def test_une_date_illisible_reste_ignoree():
    assert _en_date("la semaine prochaine") is None


def test_une_date_illisible_est_signalee_a_l_import():
    """Elle était avalée en silence : l'utilisateur croyait sa date enregistrée."""
    from io import BytesIO

    from openpyxl import Workbook

    from app.exports.excel import lire

    classeur = Workbook()
    classeur.active.append(["Entreprise", "Poste", "Date limite"])
    classeur.active.append(["Banque A", "Analyste", "la semaine prochaine"])
    tampon = BytesIO()
    classeur.save(tampon)

    reprise = lire(tampon.getvalue())
    assert reprise.lignes[0]["deadline"] is None
    assert any("illisible" in p for p in reprise.problemes)


# --- Le PDF ne peut plus être périmé ----------------------------------------


def _sans_libreoffice(monkeypatch, code_retour: int):
    """Simule un appel à LibreOffice qui n'écrit aucun fichier."""
    import subprocess

    from app.documents import pdf as module_pdf

    monkeypatch.setattr(module_pdf, "chemin_soffice", lambda: "soffice")
    monkeypatch.setattr(
        module_pdf.subprocess, "run",
        lambda *a, **kw: subprocess.CompletedProcess(a[0] if a else [], code_retour,
                                                     stdout="", stderr="verrouille"),
    )


def test_un_pdf_anterieur_a_la_conversion_est_refuse(tmp_path, monkeypatch):
    """LibreOffice déjà ouvert : la conversion n'écrit rien, mais le PDF de la
    génération précédente satisfaisait le seul contrôle d'existence."""
    from app.documents.pdf import ConversionEchouee, convertir

    import os

    source = tmp_path / "CV.docx"
    source.write_bytes(b"docx")
    perime = tmp_path / "CV.pdf"
    perime.write_bytes(b"un pdf d'hier")
    # Daté d'hier : c'est exactement l'état d'un dossier régénéré le lendemain.
    hier = perime.stat().st_mtime - 86400
    os.utime(perime, (hier, hier))

    _sans_libreoffice(monkeypatch, code_retour=0)
    with pytest.raises(ConversionEchouee, match="précédente"):
        convertir(source)


def test_un_code_de_retour_non_nul_est_signale(tmp_path, monkeypatch):
    from app.documents.pdf import ConversionEchouee, convertir

    source = tmp_path / "CV.docx"
    source.write_bytes(b"docx")

    _sans_libreoffice(monkeypatch, code_retour=1)
    with pytest.raises(ConversionEchouee, match="code 1"):
        convertir(source)


# --- Le dossier est nettoyé avant régénération ------------------------------


def _profil():
    from app.models import Profile

    return Profile(prenom="Maxime", nom="Nicolas", skills=[{"nom": "Excel"}],
                   experiences=[{"entreprise": "Crédit Mutuel", "poste": "Gestionnaire"}],
                   formations=[{"diplome": "Master", "etablissement": "EM Normandie"}])


def _offre():
    from app.models import Offer

    return Offer(source="t", source_id="1", titre="Analyste", entreprise="Banque")


@avec_modele
def test_une_lettre_refusee_ne_laisse_pas_la_precedente(tmp_path):
    """Le CV serait à jour et la lettre décrirait l'ancien profil, avec la même
    date : rien ne les distinguerait dans le dossier envoyé."""
    from app.documents.dossier import generer, nom_dossier

    offre = _offre()
    generer(_profil(), offre, tmp_path, MODELE,
            redacteur=lambda s, m: "Mon parcours au Crédit Mutuel me prépare. " * 25,
            ouvrir_apres=False)
    dossier = tmp_path / nom_dossier(offre)
    assert (dossier / "Lettre_de_motivation.docx").exists()

    resultat = generer(_profil(), offre, tmp_path, MODELE,
                       redacteur=lambda s, m: "J'ai dirigé Bridgewater dix ans. " * 25,
                       tentatives_lettre=1, ouvrir_apres=False)
    assert not (dossier / "Lettre_de_motivation.docx").exists()
    assert any("Lettre non générée" in a for a in resultat.avertissements)


@avec_modele
def test_un_fichier_personnel_depose_dans_le_dossier_est_conserve(tmp_path):
    """On n'efface que ce que l'application produit."""
    from app.documents.dossier import generer, nom_dossier

    offre = _offre()
    dossier = tmp_path / nom_dossier(offre)
    dossier.mkdir(parents=True)
    (dossier / "mes-notes.txt").write_text("préparer l'entretien", encoding="utf-8")

    generer(_profil(), offre, tmp_path, MODELE,
            redacteur=lambda s, m: "Mon parcours me prépare à ce poste. " * 25,
            ouvrir_apres=False)
    assert (dossier / "mes-notes.txt").exists()


# --- Le CV ne se contredit plus ---------------------------------------------


@avec_modele
def test_le_cv_n_affiche_plus_de_categorie_mensongere(tmp_path):
    """Le rendu portait « Quantitatif & données : R, VBA, Power BI, Word,
    PowerPoint » — Word et PowerPoint n'y ont pas leur place."""
    from app.documents.cv_render import rendre
    from app.models import Profile

    profil = Profile(prenom="M", nom="N",
                     skills=[{"nom": n} for n in ("Excel", "Word", "PowerPoint", "Python")])
    chemin = rendre(profil, _offre(), MODELE, tmp_path / "CV.docx")

    textes = [p.text for p in docx.Document(str(chemin)).paragraphs]
    for etiquette in ("Quantitatif & données :", "Finance de marché :", "Programmation :"):
        assert not any(t.startswith(etiquette) for t in textes), etiquette
    assert any("Excel" in t for t in textes), "les compétences doivent rester présentes"


@avec_modele
def test_un_profil_sans_preference_ne_produit_pas_de_tiret_orphelin(tmp_path):
    """« Recherche : — » sous le nom donne l'impression d'un document cassé."""
    from app.documents.cv_render import rendre
    from app.models import Profile

    chemin = rendre(Profile(prenom="M", nom="N", skills=[{"nom": "Excel"}]),
                    _offre(), MODELE, tmp_path / "CV.docx")
    textes = [p.text.strip() for p in docx.Document(str(chemin)).paragraphs]
    assert not any("Recherche : —" in t or t.endswith("—") for t in textes)


@avec_modele
def test_les_preferences_renseignees_apparaissent_toujours(tmp_path):
    """Les contrats recherchés figurent en en-tête — la mobilité, non.

    Ce test exigeait aussi « Mobilité : France ». La ligne a été retirée depuis :
    les pays acceptés servent à filtrer les offres, pas à figurer sur un CV. Le
    recruteur sait où est son poste, et dix-sept pays mangeaient trois lignes.
    Ce qu'il protège reste entier : une préférence renseignée doit s'afficher,
    et jamais sous forme de tiret.
    """
    from app.documents.cv_render import rendre
    from app.models import Profile

    profil = Profile(prenom="M", nom="N", skills=[{"nom": "Excel"}],
                     pays_acceptes=["France"], contrats_acceptes=["CDI"])
    chemin = rendre(profil, _offre(), MODELE, tmp_path / "CV.docx")
    textes = [p.text for p in docx.Document(str(chemin)).paragraphs]
    assert any("Recherche : CDI" in t for t in textes)
    assert not any("Mobilité" in t for t in textes)
